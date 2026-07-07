"""
建筑能耗分析 Word 报告生成器
输出格式参照《可行性报告》第三章"能耗现状分析"结构
使用 python-docx + matplotlib
"""
import sys, os, json, io, tempfile
from datetime import datetime

# Check dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    # Setup Chinese font
    plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'WenQuanYi Micro Hei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
except ImportError:
    print("ERROR: matplotlib not installed. Run: pip install matplotlib")
    sys.exit(1)


# ============================================================
# Configuration
# ============================================================
COLORS = {
    'electricity': '#f59e0b',
    'gas': '#8b5cf6',
    'heat': '#ef4444',
    'total': '#3b82f6',
    'carbon': '#10b981',
    'water': '#06b6d4',
}


def set_cell_border(cell, **kwargs):
    """Set cell border properties"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ['sz', 'val', 'color', 'space']:
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def make_table(doc, headers, rows, col_widths=None, header_style='Table Text'):
    """Create a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for ci, header in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = str(header)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = doc.styles[header_style] if header_style in [s.name for s in doc.styles] else doc.styles['Normal']
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.bold = True

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

    return table


# ============================================================
# Chart Generation
# ============================================================

def generate_charts(data, output_dir):
    """Generate chart images and return paths"""
    chart_paths = {}

    # ---- Chart 1: Energy Proportion Pie Chart ----
    proportions = data.get('energy_proportion', {}).get('proportions', [])
    if proportions:
        fig, ax = plt.subplots(figsize=(6, 5))
        labels = [p['label'] for p in proportions if p['amount'] > 0]
        values = [p['proportion_pct'] for p in proportions if p['amount'] > 0]
        colors = [COLORS['electricity'], COLORS['gas'], COLORS['heat']][:len(labels)]
        wedges, texts, autotexts = ax.pie(values, labels=labels, autopct='%1.1f%%',
                                          colors=colors, startangle=90,
                                          textprops={'fontsize': 11})
        for at in autotexts:
            at.set_fontweight('bold')
        ax.set_title('能源消耗比例拆分图', fontsize=14, fontweight='bold')
        path = os.path.join(output_dir, 'chart_pie.png')
        fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        chart_paths['pie'] = path

    # ---- Chart 2: Monthly Energy Consumption Bar Chart ----
    monthly = data.get('monthly_trend', {}).get('monthly_data', [])
    if monthly:
        fig, ax1 = plt.subplots(figsize=(10, 5))
        months = [f"{m['month']}月" for m in monthly]
        elec_vals = [m['electricity_kwh'] / 10000 for m in monthly]
        gas_vals = [m['gas_m3'] / 10000 for m in monthly]

        ax1.bar(months, elec_vals, label='电力 (万kWh)', color=COLORS['electricity'], alpha=0.85)
        ax1.bar(months, gas_vals, bottom=elec_vals, label='天然气 (万m3)', color=COLORS['gas'], alpha=0.85)
        ax1.set_ylabel('万kWh / 万m3', fontsize=11)
        ax1.set_title('逐月能源消耗走势图', fontsize=14, fontweight='bold')
        ax1.legend(loc='upper right', fontsize=9)
        ax1.set_xticks(range(len(months)))
        ax1.set_xticklabels(months, rotation=45)
        ax1.grid(axis='y', alpha=0.3)

        # Add coal line on secondary axis
        ax2 = ax1.twinx()
        coal_vals = [m['total_coal_kgce'] / 1000 for m in monthly]
        ax2.plot(months, coal_vals, 'o-', color=COLORS['total'], linewidth=2, markersize=6, label='总能耗 (tce)')
        ax2.set_ylabel('吨标准煤 (tce)', fontsize=11)
        ax2.legend(loc='upper left', fontsize=9)

        path = os.path.join(output_dir, 'chart_monthly.png')
        fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        chart_paths['monthly'] = path

    # ---- Chart 3: Carbon Emission Trend ----
    carbon_monthly = data.get('carbon_emission', {}).get('monthly_emission', [])
    if carbon_monthly:
        fig, ax = plt.subplots(figsize=(10, 4.5))
        months = [f"{m['month']}月" for m in carbon_monthly]
        co2_vals = [m['co2_total'] for m in carbon_monthly]

        ax.fill_between(range(len(months)), co2_vals, alpha=0.3, color=COLORS['carbon'])
        ax.plot(range(len(months)), co2_vals, 'o-', color=COLORS['carbon'], linewidth=2, markersize=8)
        ax.set_xticks(range(len(months)))
        ax.set_xticklabels(months, rotation=45)
        ax.set_ylabel('tCO2', fontsize=11)
        ax.set_title('逐月碳排放趋势图', fontsize=14, fontweight='bold')
        ax.grid(axis='y', alpha=0.3)

        for i, v in enumerate(co2_vals):
            ax.annotate(f'{v:.0f}', (i, v), textcoords="offset points", xytext=(0, 10),
                       ha='center', fontsize=8)

        path = os.path.join(output_dir, 'chart_carbon.png')
        fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        chart_paths['carbon'] = path

    return chart_paths


# ============================================================
# Document Assembly
# ============================================================

def _describe_energy_use(building_type):
    """Return (main_use_elec, main_use_gas) descriptions based on building type."""
    if '酒店' in building_type or '旅馆' in building_type or '宾馆' in building_type:
        return ('中央空调系统、客房与公共区域照明、电梯及办公设备',
                '锅炉房制备生活热水、供暖、厨房烹饪及洗衣房蒸汽')
    elif '学校' in building_type or '教学' in building_type or '大学' in building_type:
        return ('教学楼与办公楼照明及空调、实验室设备、数据中心、电梯及公共设施',
                '食堂厨房烹饪及热水供应')
    elif '办公' in building_type:
        return ('空调系统、办公照明、电梯、办公设备及数据中心',
                '供暖锅炉及食堂厨房')
    elif '医院' in building_type:
        return ('空调系统、医疗设备、照明、电梯及后勤设备',
                '供暖锅炉、消毒蒸汽、食堂及生活热水')
    elif '商业' in building_type or '商场' in building_type:
        return ('空调系统、公共照明、电梯扶梯及商铺用电',
                '供暖及餐饮厨房')
    else:
        return ('空调系统、照明、电梯及各类设备',
                '供暖、热水及厨房')


def generate_report(data, output_path):
    """Generate Word report adapted to building type and location."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    bi = data.get('building_info', {})
    ep = data.get('energy_proportion', {})
    mt = data.get('monthly_trend', {})
    ce = data.get('carbon_emission', {})
    sc = data.get('standard_comparison', {})

    building_name = bi.get('name', '某建筑')
    area = bi.get('area', 0)
    year = bi.get('year', 2025)
    btype = bi.get('type', '')
    location = bi.get('location', '')

    # Generate charts
    with tempfile.TemporaryDirectory() as tmpdir:
        chart_paths = generate_charts(data, tmpdir)

        # ================================================================
        h = doc.add_heading(f'{building_name}能耗分析报告', level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

        total_coal = ep.get('total_coal_kgce', 0)
        proportions = ep.get('proportions', [])
        elec_pct = gas_pct = elec_amount = gas_amount = 0
        for p in proportions:
            if p['energy_type'] == 'electricity_kwh':
                elec_pct = p['proportion_pct']
                elec_amount = p['amount']
            elif p['energy_type'] == 'gas_m3':
                gas_pct = p['proportion_pct']
                gas_amount = p['amount']

        main_elec, main_gas = _describe_energy_use(btype)
        building_label = btype + '建筑' if btype else '建筑'

        # ---- 3.1 总体能耗情况 ----
        doc.add_heading('3.1 总体能耗情况', level=2)

        doc.add_paragraph(
            f'{building_name}的能源消耗类型为电力及燃气。'
            f'电能消耗主要用于{main_elec}；'
            f'燃气消耗主要用于{main_gas}。'
        )

        # Data scope notes
        notes = data.get('data_notes', [])
        # Auto-include data validation warnings
        for w in data.get('warnings', []):
            notes.append(f'[数据告警] {w["message"]}')
        for note in notes:
            p = doc.add_paragraph(f'注：{note}')
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph(
            f'{building_name}{year}年综合能耗为{total_coal/1000:,.0f}吨标准煤（tce），'
            f'其中电力消耗{elec_amount/10000:,.1f}万kWh，折标煤{total_coal*elec_pct/100/1000:,.0f}tce，'
            f'占综合能耗的{elec_pct:.0f}%；'
            f'燃气消耗{gas_amount/10000:.2f}万m³，折标煤{total_coal*gas_pct/100/1000:,.0f}tce，'
            f'占综合能耗的{gas_pct:.0f}%。'
        )

        # ---- 3.2 建筑能耗指标及对标分析 ----
        doc.add_heading('3.2 建筑能耗指标及对标分析', level=2)

        energy_std_text = sc.get('standard_source', '国家标准')
        is_db31 = 'DB31/T 783' in energy_std_text
        doc.add_paragraph(f'{building_name}总建筑面积{area:,.0f}平方米。')
        doc.add_paragraph(f'依据{energy_std_text}对{building_label}能耗水平进行对标分析。')

        coal_per_m2 = data.get('coal_per_m2_kgce', 0)
        carbon_per_m2 = data.get('carbon_intensity_kgco2_per_m2', 0)
        ce_total = ce.get('total_emission_tons', 0)
        table_num = 1

        # ---- Table: Conversion coefficients (from the actual standard) ----
        doc.add_paragraph(f'表 3.{table_num} 能源折标准煤系数（{energy_std_text} 附录A）')
        table_num += 1
        if is_db31:
            make_table(doc,
                ['序号', '能源种类', '折标准煤系数', '备注'],
                [
                    ['1', '电力（等价值）', '0.28078 kgce/kWh', 'DB31/T 783-2026 附录A'],
                    ['2', '天然气', '1.29971 kgce/m³', 'DB31/T 783-2026 附录A'],
                ])
            doc.add_paragraph(
                '注：上述系数依据上海市统计局2024年相关规定执行。'
                'DB31/T 783-2026 附录A未规定碳排放因子，本报告碳排放计算采用生态环境部发布的区域电网排放因子。')
        else:
            grid_factor_used = ce.get('grid_factor_used', 0.581)
            gas_factor_used = ce.get('emission_breakdown', {}).get('gas', {}).get('factor', 1.997)
            region_label = ce.get('region', '全国平均')
            make_table(doc,
                ['序号', '能源种类', '折标煤系数', '碳排放因子'],
                [
                    ['1', '电力', '0.298 kgce/kWh', f'{grid_factor_used} kgCO₂/kWh'],
                    ['2', '天然气', '1.2143 kgce/m³', f'{gas_factor_used} kgCO₂/m³'],
                ])
            doc.add_paragraph(f'注：电力碳排放因子采用{region_label}电网数据。')
        doc.add_paragraph()

        # ---- Table: Standard limits ----
        if sc.get('status') == 'matched':
            has_benchmark = sc.get('energy_benchmark') is not None
            has_carbon = bool(sc.get('carbon_constraint'))

            doc.add_paragraph(f'表 3.{table_num} {btype}建筑单位建筑面积能耗限额指标 [kgce/(m²·a)]')
            table_num += 1
            if is_db31:
                scat = sc.get('school_category', '本科院校/成人高等学校')
                make_table(doc,
                    ['学校类别', '约束值', '基准值', '引导值'],
                    [[scat, str(sc.get('energy_constraint', '—')),
                      str(sc.get('energy_benchmark', '—')),
                      str(sc.get('energy_guide', '—'))]])
            elif has_benchmark:
                make_table(doc,
                    ['参考指标', '单位', '约束值', '基准值', '引导值'],
                    [['单位建筑面积能耗', 'kgce/(m²·a)',
                      f"≤{sc.get('energy_constraint', '—')}",
                      f"≤{sc.get('energy_benchmark', '—')}",
                      f"≤{sc.get('energy_guide', '—')}"]])
            else:
                make_table(doc,
                    ['参考指标', '单位', '约束值', '引导值'],
                    [['单位建筑面积能耗', 'kgce/(m²·a)',
                      f"≤{sc.get('energy_constraint', sc.get('constraint_value', '—'))}",
                      f"≤{sc.get('energy_guide', sc.get('guide_value', '—'))}"]])
            doc.add_paragraph()

            if has_carbon:
                doc.add_paragraph(f'表 3.{table_num} {btype}建筑单位建筑面积碳排放限额指标 [kgCO₂/(m²·a)]')
                table_num += 1
                make_table(doc,
                    ['参考指标', '单位', '约束值', '基准值', '引导值'],
                    [['单位建筑面积碳排放', 'kgCO₂/(m²·a)',
                      f"≤{sc.get('carbon_constraint', '—')}",
                      f"≤{sc.get('carbon_benchmark', '—')}",
                      f"≤{sc.get('carbon_guide', '—')}"]])
                doc.add_paragraph()
        else:
            doc.add_paragraph(f'（未找到适用于{building_label}的能耗限额标准，以下仅展示计算值。）')
            doc.add_paragraph()

        # ---- Correction formula (DB31/T 783 specific) ----
        if is_db31 and sc.get('correction_formula'):
            doc.add_paragraph(
                f'根据{energy_std_text}第5.3.2条，本科院校单位建筑面积年综合能耗需进行修正：')
            doc.add_paragraph(
                f'e1(修正) = E / (S × α1 × α2) = {sc.get("coal_per_m2_raw", coal_per_m2):.2f} '
                f'/ ({sc.get("alpha1", 1.0)} × {sc.get("alpha2", 1.0)}) '
                f'= {sc.get("coal_per_m2_corrected", coal_per_m2):.2f} kgce/(m²·a)')
            doc.add_paragraph(
                f'其中α1为教学科研设备资产修正系数（表3），α2为一流学科门类修正系数（表4）。')
            doc.add_paragraph()

        # ---- Calculation result table ----
        doc.add_paragraph(f'表 3.{table_num} {building_name}{year}年能耗指标计算表')
        table_num += 1
        calc_rows = [
            ['总能耗（kgce）', f'{total_coal:,.0f}'],
            ['总碳排放（kgCO₂）', f'{ce_total*1000:,.0f}'],
            ['建筑面积（m²）', f'{area:,.0f}'],
        ]
        if is_db31 and sc.get('coal_per_m2_corrected'):
            calc_rows.append(['单位面积能耗（未修正） [kgce/(m²·a)]',
                             f'{sc.get("coal_per_m2_raw", coal_per_m2):.2f}'])
            calc_rows.append(['单位面积能耗（修正后） [kgce/(m²·a)]',
                             f'{sc.get("coal_per_m2_corrected"):.2f}'])
        else:
            calc_rows.append(['单位面积能耗 [kgce/(m²·a)]', f'{coal_per_m2:.2f}'])
        calc_rows.append(['单位面积碳排放 [kgCO₂/(m²·a)]', f'{carbon_per_m2:.2f}'])
        make_table(doc, ['内容', f'{year}年'], calc_rows)
        doc.add_paragraph()

        # Standard conclusion — use clean professional text
        if sc.get('status') == 'matched':
            level = sc.get('energy_level', '') or sc.get('level', '')
            icon = sc.get('energy_icon', '') or sc.get('icon', '')
            constraint = sc.get('energy_constraint', sc.get('constraint_value', 0))
            guide = sc.get('energy_guide', sc.get('guide_value', 0))
            benchmark = sc.get('energy_benchmark')

            if level == '超标':
                verdict = (
                    f'对标结果显示，{building_name}{year}年单位建筑面积能耗为{coal_per_m2:.2f} '
                    f'kgce/(m²·a)，超出{energy_std_text}约束值（{constraint} kgce/(m²·a)），'
                    f'判定为{icon}超标。'
                )
            elif level in ('优良', '达到引导值'):
                verdict = (
                    f'对标结果显示，{building_name}{year}年单位建筑面积能耗为{coal_per_m2:.2f} '
                    f'kgce/(m²·a)，达到{energy_std_text}引导值要求，判定为{icon}{level}。'
                )
            elif level in ('达标', '达到约束值'):
                verdict = (
                    f'对标结果显示，{building_name}{year}年单位建筑面积能耗为{coal_per_m2:.2f} '
                    f'kgce/(m²·a)，满足{energy_std_text}约束值要求，判定为{icon}{level}。'
                )
            elif level == '达到基准值':
                verdict = (
                    f'对标结果显示，{building_name}{year}年单位建筑面积能耗为{coal_per_m2:.2f} '
                    f'kgce/(m²·a)，达到{energy_std_text}基准值要求（{benchmark} kgce/(m²·a)），'
                    f'距引导值（{guide} kgce/(m²·a)）尚有差距，判定为{icon}{level}。'
                )
            else:
                verdict = (
                    f'对标结果显示，{building_name}{year}年单位建筑面积能耗为{coal_per_m2:.2f} '
                    f'kgce/(m²·a)，判定为{icon}{level}。'
                )
            doc.add_paragraph(verdict)
        else:
            doc.add_paragraph(
                f'{building_name}{year}年单位建筑面积能耗为{coal_per_m2:.2f} kgce/(m²·a)，'
                f'单位建筑面积碳排放为{carbon_per_m2:.2f} kgCO₂/(m²·a)。'
            )

        # ---- 3.3 能源种类构成及占比分析 ----
        doc.add_heading('3.3 能源种类构成及占比分析', level=2)

        doc.add_paragraph(
            f'将{building_name}消耗的各类能源统一折算为标准煤，对其占比进行统计分析。'
            f'{year}年电力消耗占总能耗的{elec_pct:.0f}%，'
            f'天然气消耗占总能耗的{gas_pct:.0f}%。'
            f'电力消耗主要用于{main_elec}；'
            f'天然气消耗主要用于{main_gas}。'
        )

        doc.add_paragraph(f'表 3.5 {building_name}{year}年能源消耗统计表')
        make_table(doc,
            ['项目', '电力（kWh）', '天然气（m³）'],
            [
                ['实物量', f'{elec_amount:,.0f}', f'{gas_amount:,.0f}'],
                ['折标煤（tce）', f'{total_coal*elec_pct/100/1000:,.0f}',
                 f'{total_coal*gas_pct/100/1000:,.0f}'],
            ])
        doc.add_paragraph('注：tce——吨标准煤。电力折标煤系数依据DB31/T 783-2026附录A采用等价值0.28078 kgce/kWh。')
        doc.add_paragraph()

        if 'pie' in chart_paths:
            doc.add_paragraph(f'图 3.1 {building_name}{year}年能源消耗比例拆分图')
            doc.add_picture(chart_paths['pie'], width=Inches(4.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        # ---- 3.4 逐月能耗分析 ----
        doc.add_heading('3.4 逐月能耗分析', level=2)

        doc.add_paragraph(
            f'根据{building_name}的逐月用能数据，整理得到{year}年逐月能源消耗情况，'
            f'包括用电量（kWh）、用气量（m³）以及折算后的总能耗（tce），如下表所示。'
        )

        monthly_data = mt.get('monthly_data', [])
        if monthly_data:
            months_1_6 = [m for m in monthly_data if 1 <= m['month'] <= 6]
            months_7_12 = [m for m in monthly_data if 7 <= m['month'] <= 12]

            for half_label, half_data in [('1-6月', months_1_6), ('7-12月', months_7_12)]:
                if half_data:
                    doc.add_paragraph(
                        f'表 3.{6 if "1-6" in half_label else 7} '
                        f'{building_name}{year}年逐月能源消耗数据（{half_label}）')
                    headers = ['月份'] + [f"{m['month']}" for m in half_data]
                    make_table(doc, headers, [
                        ['电力（kWh）'] + [f"{m['electricity_kwh']:,.0f}" for m in half_data],
                        ['天然气（m³）'] + [f"{m['gas_m3']:,.0f}" for m in half_data],
                        ['折标煤（tce）'] + [f"{m['total_coal_kgce']/1000:.2f}" for m in half_data],
                    ])
                    doc.add_paragraph()

        if 'monthly' in chart_paths:
            doc.add_paragraph(f'图 3.2 {building_name}{year}年逐月用能走势图')
            doc.add_picture(chart_paths['monthly'], width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        if 'carbon' in chart_paths:
            doc.add_paragraph(f'图 3.3 {building_name}{year}年逐月碳排放趋势图')
            doc.add_picture(chart_paths['carbon'], width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        # ---- 3.5 结论 ----
        doc.add_heading('3.5 结论', level=2)

        doc.add_paragraph(
            f'通过对{building_name}{year}年能源消耗数据的分析，得出以下结论：'
        )

        # 3.5.1
        doc.add_heading('3.5.1 能源结构特征', level=3)
        if elec_pct > 90:
            struct_desc = '以电力为绝对主导，燃气消耗占比极小。'
        elif elec_pct > 60:
            struct_desc = '呈现"电力为主、燃气为辅"的特征。'
        else:
            struct_desc = '电力与燃气并重。'
        doc.add_paragraph(
            f'{building_name}的能源消耗结构{struct_desc}'
            f'{year}年电力消耗占总能耗的{elec_pct:.0f}%，'
            f'主要用于{main_elec}；'
            f'燃气消耗占{gas_pct:.0f}%，主要用于{main_gas}。'
        )

        # 3.5.2
        doc.add_heading('3.5.2 能耗季节性规律', level=3)
        peak_month = mt.get('peak_month', {})
        valley_month = mt.get('valley_month', {})
        dominant = mt.get('seasonal_analysis', {}).get('dominant_season', '')

        pk_m = peak_month.get('month', 0)
        vl_m = valley_month.get('month', 0)
        doc.add_paragraph(
            f'能耗数据呈现明显的季节性波动。'
            f'全年峰值出现在{pk_m}月，谷值出现在{vl_m}月。'
            + (f'{dominant}能耗占比最高，' if dominant else '') +
            f'夏季制冷是用电高峰的主要驱动力。'
        )

        # 3.5.3
        doc.add_heading('3.5.3 节能优化方向', level=3)
        doc.add_paragraph(
            f'基于能耗数据分析，{building_name}的节能工作可重点关注以下方面：'
        )
        doc.add_paragraph(
            '空调系统优化：夏季制冷是全年用电峰值的主因，'
            '建议对制冷机组进行能效评估，考虑变频改造、智能群控等措施。'
        )
        doc.add_paragraph(
            '用能监测精细化：建议按楼宇、分系统建立独立计量，'
            '识别高耗能单元，为节能改造和定额管理提供数据支撑。'
        )
        if gas_pct > 10:
            doc.add_paragraph(
                '燃气系统提效：燃气消耗占比较高，建议对锅炉等燃气设备进行能效检测，'
                '评估热泵替代或余热回收的可行性。'
            )

    # Save
    doc.save(output_path)
    return output_path


# ============================================================
# Main
# ============================================================
def main():
    if len(sys.argv) < 3:
        print("Usage: generate_report.py <analysis_json_file> <output_path.docx>")
        print("   or: echo '<json>' | generate_report.py - <output_path.docx>")
        sys.exit(1)

    json_input = sys.argv[1]
    output_path = sys.argv[2]

    if json_input == '-':
        raw = sys.stdin.read().strip()
    elif os.path.isfile(json_input):
        with open(json_input, 'r', encoding='utf-8') as f:
            raw = f.read()
    else:
        raw = json_input

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        print("ERROR: Invalid JSON input", file=sys.stderr)
        sys.exit(1)

    result_path = generate_report(data, output_path)
    print(f"Report saved to: {result_path}")


if __name__ == '__main__':
    main()
